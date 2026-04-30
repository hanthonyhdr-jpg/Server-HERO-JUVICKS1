import streamlit as st
import os, sys, io, base64, json, uuid, sqlite3, tempfile
from datetime import datetime, timedelta
from docx import Document

if getattr(sys, 'frozen', False):
    sys.path.append(sys._MEIPASS)
else:
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from database import buscar_dados, executar_comando
from TEMAS.moderno import apply_modern_style
from utils.auth_manager import verificar_autenticacao, obter_dados_empresa, verificar_permissao_modulo
from utils.query_cache import cached_buscar_dados, get_config_cache_version

# ══════════════════════════════════════════════════════════════════════════════
# CONVERSÃO DOCX → PDF VIA MICROSOFT WORD (Fidelidade Total ao Layout)
# ══════════════════════════════════════════════════════════════════════════════
def docx_doc_to_pdf(doc: Document) -> bytes:
    """Converte DOCX para PDF usando o Microsoft Word. Resultado 100% fiel ao arquivo."""
    tmp_docx = None
    tmp_pdf  = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_docx = tmp.name
        doc.save(tmp_docx)
        tmp_pdf = tmp_docx.replace(".docx", ".pdf")

        script = f"""
import sys, os, pythoncom
try:
    pythoncom.CoInitialize()
    from docx2pdf import convert
    convert({repr(tmp_docx)}, {repr(tmp_pdf)})
    sys.exit(0 if os.path.exists({repr(tmp_pdf)}) else 2)
except Exception as e:
    print(str(e), file=sys.stderr)
    sys.exit(1)
finally:
    try: pythoncom.CoUninitialize()
    except: pass
"""
        import subprocess
        result = subprocess.run(
            [sys.executable, "-c", script],
            timeout=120,
            capture_output=True,
            text=True
        )

        if result.returncode == 0 and os.path.exists(tmp_pdf):
            with open(tmp_pdf, "rb") as f:
                return f.read()
        else:
            err = result.stderr.strip() or f"Código {result.returncode}"
            if "Microsoft Word" in err or result.returncode == 2:
                err += "\n\n⚠️ DICA: Certifique-se de que o Microsoft Word está instalado e ativado neste computador. Ele é necessário para converter o contrato mantendo a formatação original."
            raise RuntimeError(f"Word não conseguiu converter.\nDetalhe: {err}")
    finally:
        for p in [tmp_docx, tmp_pdf]:
            if p and os.path.exists(p):
                try: os.remove(p)
                except: pass


def safe_loads(json_str, fallback=None):
    if fallback is None: fallback = []
    try: return json.loads(json_str) if json_str else fallback
    except: return fallback


# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURAÇÃO
# ══════════════════════════════════════════════════════════════════════════════
config_cache_version = get_config_cache_version()
df_conf = cached_buscar_dados("SELECT * FROM config LIMIT 1", version=config_cache_version)
nome_emp_global = df_conf.iloc[0]['empresa_nome'] if not df_conf.empty else "Sistema"
st.set_page_config(page_title=f"Contratos | {nome_emp_global}", layout="wide", page_icon="📜")
_, logo_src = obter_dados_empresa()
apply_modern_style(logo_url=logo_src, nome_empresa=nome_emp_global)
verificar_autenticacao()

if "autenticado" not in st.session_state or not st.session_state.autenticado:
    st.error("🔒 Acesso negado."); st.stop()

with st.sidebar:
    st.divider()
    st.markdown("### 🔐 Segurança de Modelos")
    senha_mestre = st.text_input("Senha Master", type="password")
    if senha_mestre == "admin123":
        st.success("Liberado ✓")
        st.session_state.pode_gerenciar = True
    else:
        st.warning("Bloqueado")


# ══════════════════════════════════════════════════════════════════════════════
# FUNÇÕES DE SUPORTE
# ══════════════════════════════════════════════════════════════════════════════
def calcular_data_entrega(prazo_dias):
    try:
        dias = int(''.join(filter(str.isdigit, str(prazo_dias))))
        data_final = datetime.now()
        while dias > 0:
            data_final += timedelta(days=1)
            if data_final.weekday() < 5: dias -= 1
        return data_final.strftime("%d/%m/%Y")
    except: return "A combinar"


def valor_por_extenso(valor):
    try:
        unidades = ["","um","dois","três","quatro","cinco","seis","sete","oito","nove"]
        dezenas  = ["","dez","vinte","trinta","quarenta","cinquenta","sessenta","setenta","oitenta","noventa"]
        especiais= ["dez","onze","doze","treze","quatorze","quinze","dezesseis","dezessete","dezoito","dezenove"]
        centenas = ["","cento","duzentos","trezentos","quatrocentos","quinhentos","seiscentos","setecentos","oitocentos","novecentos"]
        def grp(n):
            if n==100: return "cem"
            h,d,u = n//100,(n%100)//10,n%10
            r=[]
            if h: r.append(centenas[h])
            if d==1: r.append(especiais[u])
            else:
                if d>1: r.append(dezenas[d])
                if u: r.append(unidades[u])
            return " e ".join(r)
        if valor==0: return "zero reais"
        vi,vc = int(valor),int(round((valor-int(valor))*100))
        partes=[]
        if vi>=1000:
            mil,rst = vi//1000,vi%1000
            partes.append("mil" if mil==1 else f"{grp(mil)} mil")
            if rst: partes.append(("e " if rst<100 or rst%100==0 else "")+grp(rst))
        else: partes.append(grp(vi))
        res=" ".join(partes)+(" real" if vi==1 else " reais")
        if vc: res+=" e "+grp(vc)+(" centavo" if vc==1 else " centavos")
        return res.capitalize()
    except: return ""


def substituir_tags(doc, mapa):
    def processar_texto(texto):
        if not texto: return texto
        for tag, val in mapa.items():
            if tag in texto:
                texto = texto.replace(tag, str(val if val is not None else ""))
        return texto

    def processar_container(container):
        for p in container.paragraphs:
            # Estratégia robusta: verifica se a tag está no texto completo do parágrafo
            full_text = p.text
            if any(tag in full_text for tag in mapa):
                # Se encontrou, reconstrói o parágrafo mantendo o estilo do primeiro run
                new_text = processar_texto(full_text)
                if p.runs:
                    p.runs[0].text = new_text
                    for i in range(1, len(p.runs)):
                        p.runs[i].text = ""
                else:
                    p.text = new_text

    # Processa corpo, tabelas, cabeçalhos e rodapés
    processar_container(doc)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                processar_container(cell)
    for sec in doc.sections:
        processar_container(sec.header)
        processar_container(sec.footer)


ASINAME_DIR     = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assiname_app")
ASINAME_DB      = os.path.join(ASINAME_DIR, "assinaturas.db")
ASINAME_UPLOADS = os.path.join(ASINAME_DIR, "uploads")

def enviar_para_assinatura(pdf_bytes, nome_arquivo, whatsapp, cliente_nome="", razao_social="", numero_pedido=""):
    try:
        os.makedirs(ASINAME_UPLOADS, exist_ok=True)
        doc_id   = str(uuid.uuid4())
        filename = f"{doc_id}_{nome_arquivo}"
        with open(os.path.join(ASINAME_UPLOADS, filename), "wb") as f: f.write(pdf_bytes)
        conn = sqlite3.connect(ASINAME_DB)
        conn.execute("""CREATE TABLE IF NOT EXISTS documents (
            id TEXT PRIMARY KEY, original_filename TEXT, filename TEXT, whatsapp_number TEXT,
            status TEXT DEFAULT 'pending', created_at TEXT, client_name TEXT,
            razao_social TEXT, numero_pedido TEXT)""")
        conn.execute(
            "INSERT INTO documents (id,original_filename,filename,whatsapp_number,created_at,client_name,razao_social,numero_pedido,status) VALUES (?,?,?,?,?,?,?,?,?)",
            (doc_id, nome_arquivo, filename, whatsapp, datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
             cliente_nome, razao_social, numero_pedido, "pending_empresa"))
        conn.commit(); conn.close()
        return {"success": True, "doc_id": doc_id}
    except Exception as e: return {"success": False, "error": str(e)}


def gerar_dicionario_mapa(d, emp):
    agora = datetime.now()
    meses = ["janeiro","fevereiro","março","abril","maio","junho","julho","agosto","setembro","outubro","novembro","dezembro"]
    prazo = str(d.get('prazo_entrega', ''))
    prazo = prazo.split(" (")[0] if " (" in prazo else prazo
    itens = safe_loads(d.get('itens_json', '[]'))
    txt_it = ""
    for i, it in enumerate(itens, 1):
        txt_it += f"{i}. {it.get('produto','')} - Qtd: {it.get('qtd','')} {it.get('unidade','')} - R$ {float(it.get('subtotal',0)):,.2f}\n"
        if it.get('descricao'): txt_it += f"   Obs: {it['descricao']}\n"
    return {
        "{{numero_orcamento}}": d.get('id',''),   "{{id_orcamento}}":     d.get('id',''),
        "{{cliente_nome}}":     str(d.get('c_nome','')).upper(),
        "[NOME DO CLIENTE]":    str(d.get('c_nome','')).upper(),
        "{{cliente_cnpj}}":     d.get('c_cnpj',''),  "[CPF/CNPJ DO CLIENTE]": d.get('c_cnpj',''),
        "{{cliente_endereco}}": d.get('c_end',''),   "{{cliente_numero}}":   d.get('c_num',''),
        "{{cliente_bairro}}":   d.get('c_bai',''),
        "{{cliente_cidade}}":   d.get('c_cid',''),   "{{cliente_uf}}":       d.get('c_uf',''),
        "{{cliente_cep}}":      d.get('c_cep',''),   "{{cliente_telefone}}": d.get('c_tel',''),
        "{{cliente_email}}":    d.get('c_email',''),
        "{{valor_total}}":      f"R$ {d.get('total',0):,.2f}",
        "{{valor_extenso}}":    valor_por_extenso(float(d.get('total',0))),
        "{{forma_pagamento}}":  d.get('forma_pagamento',''),
        "{{financeiro_obs}}":   str(d.get('observacoes','') or d.get('obs','') or d.get('condicoes_pagamento','') or ""),
        "[FINANCEIRO_OBS]":     str(d.get('observacoes','') or d.get('obs','') or d.get('condicoes_pagamento','') or ""),
        "{{observacoes}}":      str(d.get('observacoes','') or d.get('obs','') or ""),
        "{{prazo_entrega}}":    prazo,
        "{{itens_orcamento}}":  txt_it,
        "{{descricao_itens}}":  txt_it, # Tag alternativa
        "{{vendedor}}":         d.get('vendedor',''),
        "{{empresa_nome}}":     str(emp.get('empresa_nome','')).upper(),
        "[NOME DA SUA EMPRESA]":str(emp.get('empresa_nome','')).upper(),
        "{{empresa_cnpj}}":     emp.get('empresa_cnpj',''),  "[SEU CNPJ]": emp.get('empresa_cnpj',''),
        "{{empresa_endereco}}": emp.get('empresa_end',''),
        "{{empresa_numero}}":   emp.get('empresa_num',''),
        "{{empresa_telefone}}": emp.get('empresa_tel',''),
        "{{dia}}": agora.day, "{{mes}}": meses[agora.month-1], "{{ano}}": agora.year,
        "{{data_por_extenso}}": f"{agora.day} de {meses[agora.month-1]} de {agora.year}",
    }


def _preparar_doc(mod_data, d, emp):
    """Abre o modelo e substitui as tags. Retorna Document pronto."""
    doc = Document(io.BytesIO(mod_data['arquivo_bin']))
    substituir_tags(doc, gerar_dicionario_mapa(d, emp))
    return doc


def _get_pdf(cache_key, mod_data, d, emp, mod_sel_key, cur_mod):
    """Gera PDF via Word e armazena no cache. Retorna bytes."""
    if cache_key not in st.session_state or st.session_state.get(mod_sel_key) != cur_mod:
        doc = _preparar_doc(mod_data, d, emp)
        st.session_state[cache_key]   = docx_doc_to_pdf(doc)
        st.session_state[mod_sel_key] = cur_mod
    return st.session_state[cache_key]


# ══════════════════════════════════════════════════════════════════════════════
# INTERFACE
# ══════════════════════════════════════════════════════════════════════════════
st.title("📜 Gestão de Contratos Jurídicos")
tab_aprovados, tab_avulso = st.tabs(["✅ Orçamentos Aprovados", "✨ Gerar Avulso"])

# ── TAB 1: APROVADOS ──────────────────────────────────────────────────────────
with tab_aprovados:
    # Busca orçamentos aprovados com dados do cliente (JOIN) para preenchimento de tags
    df_aprov = cached_buscar_dados("""
        SELECT 
            o.id, o.cliente_id, o.vendedor, o.total, o.forma_pagamento, o.data, o.itens_json, o.obs_geral, o.prazo_entrega,
            c.nome as c_nome, c.cnpj as c_cnpj, c.telefone as c_tel, c.endereco as c_end, 
            c.numero as c_num, c.bairro as c_bai, c.cidade as c_cid, c.estado as c_uf, c.cep as c_cep, c.email as c_email,
            c.nome_fantasia as c_fantasia
        FROM orcamentos o
        LEFT JOIN clientes c ON o.cliente_id = c.id
        WHERE o.status='Aprovado'
        ORDER BY o.id DESC
    """, version=config_cache_version)

    
    df_mod = cached_buscar_dados("SELECT * FROM modelos_contrato", version=config_cache_version)
    modelos_nomes = df_mod['nome'].tolist() if not df_mod.empty else []

    if df_aprov.empty:
        st.info("💡 Nenhum orçamento aprovado ainda.")
    else:
        emp_d = df_conf.iloc[0].to_dict() if not df_conf.empty else {}

        for _, d in df_aprov.iterrows():
            oid           = d['id']
            cache_pdf_key = f"pdf_{oid}"
            mod_key       = f"mod_{oid}"
            show_prev_key = f"prev_{oid}"

            # ── Card Premium ──
            st.markdown(f"""
                <div style="background:rgba(15,40,71,0.4);border:1px solid rgba(255,255,255,0.08);
                            border-radius:16px 16px 0 0;padding:0;margin-bottom:-15px;margin-top:22px;">
                    <div style="background:linear-gradient(90deg,rgba(99,102,241,0.12),transparent);
                                padding:12px 20px;border-bottom:1px solid rgba(255,255,255,0.05);
                                display:flex;justify-content:space-between;align-items:center;">
                        <span style="font-weight:700;color:#a5b4fc;letter-spacing:1px;">🤝 CONTRATO → REF: {oid}</span>
                        <span style="background:rgba(34,197,94,0.15);border:1px solid #22c55e;color:#22c55e;
                                     padding:2px 12px;border-radius:20px;font-size:0.65rem;font-weight:800;">✅ APROVADO</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)

            with st.container(border=True):
                ci, cm, ca = st.columns([1.5, 1.3, 1.8])

                # Informações do contrato
                with ci:
                    def robust_parse_date(d_str):
                        for fmt in ('%Y-%m-%d', '%d/%m/%Y'):
                            try:
                                return datetime.strptime(str(d_str), fmt)
                            except:
                                continue
                        return datetime.now()

                    data_fmt = robust_parse_date(d['data']).strftime('%d/%m/%Y')
                    st.markdown(f"""
                        <div style='margin-bottom:10px;'>
                            <span style='color:#94a3b8;font-size:.72rem;font-weight:600;text-transform:uppercase;'>👤 Cliente</span><br>
                            <span style='font-size:1.05rem;font-weight:700;color:#f8fafc;'>{d['c_nome']}</span>
                        </div>
                        <div>
                            <span style='color:#94a3b8;font-size:.72rem;font-weight:600;text-transform:uppercase;'>💰 Total &nbsp;·&nbsp; 📅 Data</span><br>
                            <span style='font-size:1.05rem;font-weight:800;color:#22c55e;'>R$ {d['total']:,.2f}</span>
                            <span style='color:#94a3b8;'> · {data_fmt}</span>
                        </div>
                    """, unsafe_allow_html=True)

                # Seletor de modelo + botão prévia
                with cm:
                    if not modelos_nomes:
                        st.warning("Nenhum modelo .docx cadastrado.")
                        sel_mod = None
                    else:
                        sel_mod = st.selectbox("📋 Modelo:", modelos_nomes, key=mod_key)
                        st.write("")
                        if st.button("👁️ VER PRÉVIA", key=f"btn_prev_{oid}", use_container_width=True):
                            st.session_state[show_prev_key] = not st.session_state.get(show_prev_key, False)

                # Ações — apenas PDF (Word) + Assinatura
                with ca:
                    st.markdown("<div style='font-size:.7rem;font-weight:700;color:#94a3b8;letter-spacing:1px;margin-bottom:8px;'>📁 GERAR CONTRATO</div>", unsafe_allow_html=True)

                    if sel_mod:
                        mod_data = df_mod[df_mod['nome'] == sel_mod].iloc[0]

                        # Botão principal: gerar PDF via Word
                        if st.button("📄 GERAR PDF (Word)", key=f"btn_pf_{oid}", use_container_width=True, type="primary"):
                            with st.spinner("⏳ Convertendo via Word... aguarde."):
                                try:
                                    _get_pdf(cache_pdf_key, mod_data, d, emp_d, f"lm_{oid}", sel_mod)
                                    st.success("✅ PDF gerado com sucesso!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Erro: {e}")

                        # Botão de download — aparece logo após gerar
                        if cache_pdf_key in st.session_state:
                            st.download_button(
                                "🔽 Baixar PDF",
                                st.session_state[cache_pdf_key],
                                f"Contrato_{oid}.pdf",
                                mime="application/pdf",
                                key=f"dl_pf_{oid}",
                                use_container_width=True,
                                type="primary"
                            )

                        # Assinatura Digital
                        st.markdown("<hr style='margin:10px 0;opacity:.1;'>", unsafe_allow_html=True)
                        wpp = str(d.get('c_tel','') or '').replace(' ','').replace('-','').replace('(','').replace(')','')
                        if not wpp.startswith('55') and wpp: wpp = '55' + wpp.lstrip('+')

                        if st.button("✍️ ENVIAR PARA ASSINATURA", key=f"btn_sig_{oid}", use_container_width=True, type="secondary"):
                            with st.spinner("⏳ Preparando e enviando..."):
                                try:
                                    pdf_b = _get_pdf(cache_pdf_key, mod_data, d, emp_d, f"lm_{oid}", sel_mod)
                                    res   = enviar_para_assinatura(
                                        pdf_b, f"Contrato_{oid}.pdf",
                                        wpp or '5500000000000',
                                        d.get('c_fantasia') or d['c_nome'],
                                        d['c_nome'], str(oid))
                                    if res["success"]:
                                        st.success("✅ Enviado para assinatura!")
                                        st.page_link("pages/8_8-Assinaturas.py", label="→ Ir para Assinaturas", icon="✍️")
                                    else:
                                        st.error(f"Erro: {res.get('error')}")
                                except Exception as e:
                                    st.error(f"Erro na conversão: {e}")

            # ── Prévia: usa o PDF do cache (mesmo arquivo!) ──────────────
            if st.session_state.get(show_prev_key, False):
                if sel_mod and not df_mod.empty:
                    mod_data_p = df_mod[df_mod['nome'] == sel_mod].iloc[0]
                    if cache_pdf_key not in st.session_state:
                        with st.spinner("⏳ Gerando prévia via Word..."):
                            try:
                                _get_pdf(cache_pdf_key, mod_data_p, d, emp_d, f"lm_{oid}", sel_mod)
                            except Exception as e:
                                st.error(f"Erro: {e}")
                                st.session_state[show_prev_key] = False

                    if cache_pdf_key in st.session_state:
                        b64 = base64.b64encode(st.session_state[cache_pdf_key]).decode()
                        st.markdown(
                            f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="700" '
                            f'style="border:none;border-radius:10px;margin-top:12px;"></iframe>',
                            unsafe_allow_html=True)
                        if st.button("❌ Fechar Prévia", key=f"btn_cls_prev_{oid}", use_container_width=True):
                            st.session_state[show_prev_key] = False; st.rerun()


# ── TAB 2: AVULSO ─────────────────────────────────────────────────────────────
with tab_avulso:
    df_orc_av = cached_buscar_dados("""
        SELECT 
            o.id, o.cliente_id, o.vendedor, o.total, o.forma_pagamento, o.data, o.itens_json, o.obs_geral, o.prazo_entrega,
            c.nome as c_nome, c.cnpj as c_cnpj, c.telefone as c_tel, c.endereco as c_end, 
            c.numero as c_num, c.bairro as c_bai, c.cidade as c_cid, c.estado as c_uf, c.cep as c_cep, c.email as c_email,
            c.nome_fantasia as c_fantasia
        FROM orcamentos o
        LEFT JOIN clientes c ON o.cliente_id = c.id
        WHERE o.status='Aprovado'
        ORDER BY o.id DESC
    """, version=config_cache_version)
    
    df_mod_av = cached_buscar_dados("SELECT * FROM modelos_contrato", version=config_cache_version)

    if df_orc_av.empty:
        st.info("💡 Realize uma venda primeiro para gerar contratos.")
    elif df_mod_av.empty:
        st.warning("Nenhum modelo .docx cadastrado.")
    else:
        emp_d = df_conf.iloc[0].to_dict() if not df_conf.empty else {}
        c1, c2 = st.columns([1, 2])

        with c1:
            st.subheader("📋 Parâmetros")
            sel_orc_str = st.selectbox(
                "Venda de Referência:",
                [f"{r['id']} - {r['c_nome']}" for _, r in df_orc_av.iterrows()],
                key="av_orc")
            id_sel    = sel_orc_str.split(" - ")[0]
            d_av      = df_orc_av[df_orc_av['id'] == id_sel].iloc[0]
            sel_mod_av= st.selectbox("Modelo:", df_mod_av['nome'].tolist(), key="av_mod")
            mod_av    = df_mod_av[df_mod_av['nome'] == sel_mod_av].iloc[0]
            cache_av  = f"av_pdf_{id_sel}_{sel_mod_av}"

            st.write("")
            if st.button("📄 GERAR PDF (Word)", type="primary", use_container_width=True, key="av_btn_pdf"):
                with st.spinner("⏳ Convertendo via Word..."):
                    try:
                        _get_pdf(cache_av, mod_av, d_av, emp_d, "av_lm", sel_mod_av)
                        st.success("PDF pronto!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {e}")

            if cache_av in st.session_state:
                st.download_button(
                    "🔽 Baixar PDF", st.session_state[cache_av],
                    f"Contrato_{id_sel}.pdf", mime="application/pdf",
                    use_container_width=True, key="av_dl_pdf", type="primary")

            st.write("---")
            wpp_av = str(d_av.get('c_tel','') or '').replace(' ','').replace('-','').replace('(','').replace(')','')
            if not wpp_av.startswith('55') and wpp_av: wpp_av = '55' + wpp_av.lstrip('+')

            if st.button("✍️ Enviar p/ Assinatura", use_container_width=True, key="av_sig"):
                with st.spinner("⏳ Preparando..."):
                    try:
                        pdf_b = _get_pdf(cache_av, mod_av, d_av, emp_d, "av_lm", sel_mod_av)
                        res   = enviar_para_assinatura(
                            pdf_b, f"Contrato_{id_sel}.pdf",
                            wpp_av or '5500000000000',
                            d_av.get('c_fantasia') or d_av['c_nome'], d_av['c_nome'], str(id_sel))
                        if res["success"]:
                            st.success("✅ Enviado!")
                            st.page_link("pages/8_8-Assinaturas.py", label="Ver Assinaturas", icon="✍️")
                        else:
                            st.error(f"Erro: {res.get('error')}")
                    except Exception as e:
                        st.error(f"Erro: {e}")

        with c2:
            st.subheader("👁️ Prévia do Contrato")
            if st.button("🔄 Gerar / Atualizar Prévia", use_container_width=True, key="av_prev"):
                with st.spinner("⏳ Gerando prévia via Word..."):
                    try:
                        _get_pdf(cache_av, mod_av, d_av, emp_d, "av_lm", sel_mod_av)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro: {e}")

            if cache_av in st.session_state:
                b64 = base64.b64encode(st.session_state[cache_av]).decode()
                st.markdown(
                    f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="700" '
                    f'style="border:none;border-radius:10px;"></iframe>',
                    unsafe_allow_html=True)
            else:
                st.info("Clique em 'Gerar / Atualizar Prévia' para visualizar.")


